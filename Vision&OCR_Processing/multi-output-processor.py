# output/multi_output_processor.py
import os
import json
import base64
import logging
import requests
from PIL import Image
import numpy as np
import cv2
from typing import Dict, Any, List, Optional, Union, BinaryIO, Callable
from datetime import datetime
import io
import tempfile

class MultiOutputProcessor:
    """Multi-Output-Prozessor für das Universal OCR Tool 2.0.
    
    Ermöglicht die Generierung verschiedener Ausgabeformate aus einem einzigen
    Eingabedokument. Unterstützt verschiedene Ausgabetypen wie API-Weiterleitungen,
    strukturierte Daten, Bilder mit Annotationen und kombinierte Dokumente.
    
    Features:
    - Multiple Ausgabeformate aus einem Prozessierungsschritt
    - Flexible Konfiguration von Ausgabeformaten
    - Unterstützung für API-Integration mit verschiedenen Zielsystemen
    - Bildannotation und Thumbnail-Generierung
    - Kombinierte Ausgabedokumente mit Text und Bildern
    """
    
    def __init__(self, config: Dict[str, Any]):
        """Initialisiert den Multi-Output-Prozessor mit Konfiguration.
        
        Args:
            config: Konfigurationswörterbuch mit Ausgabeoptionen und Einstellungen
        """
        self.config = config
        self.logger = logging.getLogger("ocrtool.output")
        
        # Konfigurationsoptionen
        self.output_dir = config.get("output_dir", "output")
        self.enable_api_forwarding = config.get("enable_api_forwarding", False)
        self.api_endpoints = config.get("api_endpoints", [])
        self.default_formats = config.get("default_formats", ["text", "json"])
        self.enable_thumbnails = config.get("enable_thumbnails", True)
        self.thumbnail_size = config.get("thumbnail_size", (200, 200))
        self.max_filename_length = config.get("max_filename_length", 100)
        
        # Erstelle Ausgabeverzeichnis falls nicht vorhanden
        os.makedirs(self.output_dir, exist_ok=True)
        
        # Registrierte Ausgabeformatprozessoren
        self.output_processors = {
            "text": self._generate_text_output,
            "json": self._generate_json_output,
            "pdf": self._generate_pdf_output,
            "image_annotated": self._generate_annotated_image,
            "thumbnail": self._generate_thumbnail,
            "api": self._send_to_api,
            "csv": self._generate_csv_output,
            "markdown": self._generate_markdown_output,
            "html": self._generate_html_output,
            "docx": self._generate_docx_output,
            "combined": self._generate_combined_output
        }
        
        self.logger.info("Multi-Output-Prozessor initialisiert")
    
    def process(self, content: Dict[str, Any], formats: Optional[List[str]] = None,
                options: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Verarbeitet Inhalte in verschiedene Ausgabeformate.
        
        Args:
            content: Eingabeinhalt mit erkanntem Text, Metadaten und optional Bildern
            formats: Liste der gewünschten Ausgabeformate; falls None, werden Standardformate verwendet
            options: Zusätzliche verarbeitungsspezifische Optionen
            
        Returns:
            Dictionary mit Verarbeitungsergebnissen und Dateipfaden
        """
        if formats is None:
            formats = self.default_formats
        
        if options is None:
            options = {}
        
        # Basisname für Ausgabedateien
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_id = content.get("doc_id", "doc")
        base_filename = f"{doc_id}_{timestamp}"
        base_filename = self._sanitize_filename(base_filename)
        
        # Ergebnisse für verschiedene Formate
        results = {
            "base_info": {
                "timestamp": timestamp,
                "doc_id": doc_id,
                "formats_processed": []
            },
            "outputs": {}
        }
        
        # Verarbeite jedes angeforderte Format
        for output_format in formats:
            if output_format in self.output_processors:
                try:
                    # Generiere einen formatspezifischen Dateinamen
                    format_filename = f"{base_filename}.{self._get_extension(output_format)}"
                    format_path = os.path.join(self.output_dir, format_filename)
                    
                    # Verarbeite Format mit entsprechendem Prozessor
                    format_result = self.output_processors[output_format](
                        content, format_path, options
                    )
                    
                    # Speichere Ergebnis
                    if format_result:
                        results["outputs"][output_format] = format_result
                        results["base_info"]["formats_processed"].append(output_format)
                except Exception as e:
                    self.logger.error(f"Fehler bei Verarbeitung von Format {output_format}: {e}")
                    results["outputs"][output_format] = {"error": str(e)}
            else:
                self.logger.warning(f"Unbekanntes Ausgabeformat: {output_format}")
                results["outputs"][output_format] = {"error": "Format nicht unterstützt"}
        
        return results
    
    def _sanitize_filename(self, filename: str) -> str:
        """Bereinigt einen Dateinamen von ungültigen Zeichen.
        
        Args:
            filename: Ursprünglicher Dateiname
            
        Returns:
            Bereinigter Dateiname
        """
        # Entferne ungültige Zeichen
        valid_chars = "-_.() abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789"
        sanitized = ''.join(c for c in filename if c in valid_chars)
        
        # Begrenze die Länge
        if len(sanitized) > self.max_filename_length:
            name_part, ext_part = os.path.splitext(sanitized)
            max_name_length = self.max_filename_length - len(ext_part)
            sanitized = name_part[:max_name_length] + ext_part
        
        return sanitized
    
    def _get_extension(self, output_format: str) -> str:
        """Bestimmt die Dateierweiterung für ein Ausgabeformat.
        
        Args:
            output_format: Ausgabeformat
            
        Returns:
            Passende Dateierweiterung
        """
        extensions = {
            "text": "txt",
            "json": "json",
            "pdf": "pdf",
            "image_annotated": "png",
            "thumbnail": "jpg",
            "csv": "csv",
            "markdown": "md",
            "html": "html",
            "docx": "docx",
            "combined": "pdf"
        }
        
        return extensions.get(output_format, "dat")
    
    def _generate_text_output(self, content: Dict[str, Any], output_path: str,
                             options: Dict[str, Any]) -> Dict[str, Any]:
        """Generiert einfache Textausgabe.
        
        Args:
            content: Verarbeiteter Inhalt mit erkanntem Text
            output_path: Pfad für die Ausgabedatei
            options: Verarbeitungsoptionen
            
        Returns:
            Ergebnisinformationen
        """
        text = content.get("text", "")
        
        # Anwenden von Textformatierungsoptionen
        if options.get("include_confidence", False) and "text_blocks" in content:
            formatted_text = []
            for block in content["text_blocks"]:
                block_text = block.get("text", "")
                confidence = block.get("confidence", 0.0)
                formatted_text.append(f"{block_text} [Konfidenz: {confidence:.2f}]")
            text = "\n\n".join(formatted_text)
        
        # Speichern der Textdatei
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        return {
            "path": output_path,
            "format": "text",
            "size": len(text),
            "char_count": len(text)
        }
    
    def _generate_json_output(self, content: Dict[str, Any], output_path: str,
                             options: Dict[str, Any]) -> Dict[str, Any]:
        """Generiert strukturierte JSON-Ausgabe.
        
        Args:
            content: Verarbeiteter Inhalt
            output_path: Pfad für die Ausgabedatei
            options: Verarbeitungsoptionen
            
        Returns:
            Ergebnisinformationen
        """
        # Bereite JSON-Ausgabe vor
        output_data = {}
        
        # Füge grundlegende Informationen hinzu
        output_data["document_id"] = content.get("doc_id", "")
        output_data["timestamp"] = datetime.now().isoformat()
        output_data["text"] = content.get("text", "")
        
        # Füge strukturierte Daten hinzu, wenn verfügbar
        if "text_blocks" in content:
            output_data["text_blocks"] = content["text_blocks"]
        
        if "metadata" in content:
            output_data["metadata"] = content["metadata"]
        
        # Füge Erkennungsdetails hinzu, wenn verfügbar und gewünscht
        if "detections" in content and options.get("include_detections", True):
            output_data["detections"] = content["detections"]
        
        # Je nach Optionen Bilder einbetten oder ausschließen
        if "images" in content and options.get("include_images", False):
            # Konvertiere Bilder zu Base64
            output_data["images"] = {}
            for image_key, image_data in content["images"].items():
                if isinstance(image_data, np.ndarray):
                    # Konvertiere NumPy-Array zu Base64
                    _, buffer = cv2.imencode(".png", image_data)
                    output_data["images"][image_key] = base64.b64encode(buffer).decode('utf-8')
                elif isinstance(image_data, str) and os.path.exists(image_data):
                    # Lade Bilddatei und konvertiere zu Base64
                    with open(image_data, "rb") as img_file:
                        output_data["images"][image_key] = base64.b64encode(img_file.read()).decode('utf-8')
        
        # Speichern der JSON-Datei
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, ensure_ascii=False, indent=2)
        
        return {
            "path": output_path,
            "format": "json",
            "size": os.path.getsize(output_path),
            "keys": list(output_data.keys())
        }
    
    def _generate_pdf_output(self, content: Dict[str, Any], output_path: str,
                            options: Dict[str, Any]) -> Dict[str, Any]:
        """Generiert PDF-Ausgabe.
        
        Args:
            content: Verarbeiteter Inhalt
            output_path: Pfad für die Ausgabedatei
            options: Verarbeitungsoptionen
            
        Returns:
            Ergebnisinformationen
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
        except ImportError:
            self.logger.error("ReportLab nicht installiert. PDF-Generierung nicht möglich.")
            raise ImportError("ReportLab wird für PDF-Generierung benötigt. Installieren Sie 'reportlab'.")
        
        # Erstelle PDF-Dokument
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Benutzerdefinierter Stil für erkannten Text
        styles.add(ParagraphStyle(
            name='OCRText',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
        ))
        
        # Elemente für das PDF
        elements = []
        
        # Titel und Metadaten
        title = content.get("doc_id", "OCR-Ergebnis")
        elements.append(Paragraph(f"<b>{title}</b>", styles["Title"]))
        elements.append(Spacer(1, 0.25*inch))
        
        # Zeitstempel
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        elements.append(Paragraph(f"Generiert am: {timestamp}", styles["Normal"]))
        elements.append(Spacer(1, 0.25*inch))
        
        # Füge Bild hinzu, wenn vorhanden und gewünscht
        if "images" in content and options.get("include_images", True):
            original_image = content["images"].get("original")
            if original_image is not None:
                # Temporäre Bilddatei für ReportLab erstellen
                if isinstance(original_image, np.ndarray):
                    with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                        temp_filename = temp_file.name
                        cv2.imwrite(temp_filename, original_image)
                        
                    img = RLImage(temp_filename, width=6*inch, height=4*inch)
                    elements.append(img)
                    elements.append(Spacer(1, 0.25*inch))
                    
                    # Lösche temporäre Datei
                    try:
                        os.unlink(temp_filename)
                    except:
                        pass
                elif isinstance(original_image, str) and os.path.exists(original_image):
                    img = RLImage(original_image, width=6*inch, height=4*inch)
                    elements.append(img)
                    elements.append(Spacer(1, 0.25*inch))
        
        # Füge erkannten Text hinzu
        if "text" in content and content["text"]:
            elements.append(Paragraph("<b>Erkannter Text:</b>", styles["Heading2"]))
            elements.append(Spacer(1, 0.1*inch))
            
            text = content["text"]
            # Teile Text in Absätze
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    elements.append(Paragraph(para.replace('\n', '<br/>'), styles["OCRText"]))
                    elements.append(Spacer(1, 0.1*inch))
        
        # Erstelle das PDF
        doc.build(elements)
        
        return {
            "path": output_path,
            "format": "pdf",
            "size": os.path.getsize(output_path),
            "pages": 1  # In einer erweiterten Version könnte die tatsächliche Seitenzahl berechnet werden
        }
    
    def _generate_annotated_image(self, content: Dict[str, Any], output_path: str,
                                options: Dict[str, Any]) -> Dict[str, Any]:
        """Generiert ein annotiertes Bild mit erkannten Textregionen und Objekten.
        
        Args:
            content: Verarbeiteter Inhalt
            output_path: Pfad für die Ausgabedatei
            options: Verarbeitungsoptionen
            
        Returns:
            Ergebnisinformationen
        """
        # Überprüfe, ob ein Bild vorhanden ist
        if "images" not in content or "original" not in content["images"]:
            self.logger.error("Kein Bild für Annotation vorhanden")
            raise ValueError("Kein Bild für Annotation verfügbar")
        
        # Hole Originalbild
        original_image = content["images"]["original"]
        
        if isinstance(original_image, str) and os.path.exists(original_image):
            # Lade Bild aus Dateipfad
            image = cv2.imread(original_image)
        elif isinstance(original_image, np.ndarray):
            # Verwende NumPy-Array direkt
            image = original_image.copy()
        else:
            self.logger.error("Ungültiges Bildformat für Annotation")
            raise ValueError("Bild muss als NumPy-Array oder gültiger Dateipfad vorliegen")
        
        # Hole erkannte Textblöcke
        annotated_image = image.copy()
        
        # Zeichne Textregionen ein, falls vorhanden
        if "text_blocks" in content:
            for i, block in enumerate(content["text_blocks"]):
                if "bbox" in block:
                    x, y, w, h = block["bbox"]
                    confidence = block.get("confidence", 0.0)
                    
                    # Farbabstufung basierend auf Konfidenz (Rot bis Grün)
                    color = (0, int(255 * confidence), int(255 * (1 - confidence)))
                    
                    # Zeichne Begrenzungsrahmen
                    cv2.rectangle(annotated_image, (x, y), (x + w, y + h), color, 2)
                    
                    # Zeige Blocknummer und Konfidenz
                    if options.get("show_confidence", True):
                        label = f"#{i+1}: {confidence:.2f}"
                        cv2.putText(annotated_image, label, (x, y - 5),
                                  cv2.FONT_HERSHEY_SIMPLEX, 0.5, color, 2)
        
        # Zeichne erkannte Objekte ein, falls vorhanden
        if "detections" in content and options.get("show_detections", True):
            for detection in content["detections"]:
                if "box" in detection:
                    box = detection["box"]
                    x, y = box["x"], box["y"]
                    w, h = box["width"], box["height"]
                    class_name = detection.get("class_name", "Objekt")
                    confidence = detection.get("confidence", 1.0)
                    
                    # Zeichne Objektrahmen mit anderer Farbe als Textblöcke
                    cv2.rectangle(annotated_image, (x, y), (x + w, y + h), (255, 165, 0), 2)
                    
                    # Zeige Klassenname und Konfidenz
                    label = f"{class_name}: {confidence:.2f}"
                    cv2.putText(annotated_image, label, (x, y - 5),
                               cv2.FONT_HERSHEY_SIMPLEX, 0.5, (255, 165, 0), 2)
        
        # Füge Zeitstempel und Überschrift hinzu, falls gewünscht
        if options.get("add_timestamp", True):
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            cv2.putText(annotated_image, timestamp, (10, 30),
                       cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 0, 255), 2)
        
        # Speichere annotiertes Bild
        cv2.imwrite(output_path, annotated_image)
        
        return {
            "path": output_path,
            "format": "image_annotated",
            "size": os.path.getsize(output_path),
            "dimensions": f"{annotated_image.shape[1]}x{annotated_image.shape[0]}"
        }
    
    def _generate_thumbnail(self, content: Dict[str, Any], output_path: str,
                          options: Dict[str, Any]) -> Dict[str, Any]:
        """Generiert ein Thumbnail des Eingabebildes.
        
        Args:
            content: Verarbeiteter Inhalt
            output_path: Pfad für die Ausgabedatei
            options: Verarbeitungsoptionen
            
        Returns:
            Ergebnisinformationen
        """
        # Überprüfe, ob ein Bild vorhanden ist
        if "images" not in content or "original" not in content["images"]:
            self.logger.error("Kein Bild für Thumbnail-Generierung vorhanden")
            raise ValueError("Kein Bild für Thumbnail verfügbar")
        
        # Hole Originalbild
        original_image = content["images"]["original"]
        
        if isinstance(original_image, str) and os.path.exists(original_image):
            # Lade Bild mit PIL für bessere Größenänderung
            image = Image.open(original_image)
        elif isinstance(original_image, np.ndarray):
            # Konvertiere NumPy-Array zu PIL
            image = Image.fromarray(cv2.cvtColor(original_image, cv2.COLOR_BGR2RGB))
        else:
            self.logger.error("Ungültiges Bildformat für Thumbnail")
            raise ValueError("Bild muss als NumPy-Array oder gültiger Dateipfad vorliegen")
        
        # Bestimme Thumbnail-Größe
        thumbnail_size = options.get("thumbnail_size", self.thumbnail_size)
        
        # Generiere Thumbnail mit Beibehaltung des Seitenverhältnisses
        image.thumbnail(thumbnail_size, Image.LANCZOS)
        
        # Optional Text hinzufügen
        if options.get("add_text", False) and "text" in content:
            # Konvertiere zurück zu OpenCV für Textüberlagerung
            thumbnail = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            
            # Extrahiere die ersten N Zeichen des Textes
            text_preview = content["text"][:50]
            if len(content["text"]) > 50:
                text_preview += "..."
            
            # Füge Text hinzu
            cv2.putText(thumbnail, text_preview, (5, image.height - 10),
                      cv2.FONT_HERSHEY_SIMPLEX, 0.3, (255, 255, 255), 1)
            
            # Speichere als OpenCV-Bild
            cv2.imwrite(output_path, thumbnail)
        else:
            # Speichere als PIL-Bild
            image.save(output_path, quality=85)
        
        return {
            "path": output_path,
            "format": "thumbnail",
            "size": os.path.getsize(output_path),
            "dimensions": f"{image.width}x{image.height}"
        }
    
    def _send_to_api(self, content: Dict[str, Any], output_path: str,
                   options: Dict[str, Any]) -> Dict[str, Any]:
        """Sendet Daten an eine externe API.
        
        Args:
            content: Verarbeiteter Inhalt
            output_path: Wird hier nur für konsistente Schnittstelle verwendet
            options: Verarbeitungsoptionen mit API-Konfiguration
            
        Returns:
            Ergebnisinformationen mit API-Antwort
        """
        if not self.enable_api_forwarding:
            self.logger.warning("API-Weiterleitung ist deaktiviert")
            return {"error": "API-Weiterleitung deaktiviert"}
        
        # API-Endpunkt aus Optionen oder Konfiguration
        api_endpoint = options.get("api_endpoint")
        if not api_endpoint:
            # Verwende den ersten konfigurierten Endpunkt
            if not self.api_endpoints:
                self.logger.error("Kein API-Endpunkt konfiguriert")
                return {"error": "Kein API-Endpunkt verfügbar"}
            api_endpoint = self.api_endpoints[0]
        
        # Bereite Daten vor
        payload = {
            "document_id": content.get("doc_id", ""),
            "timestamp": datetime.now().isoformat(),
            "text": content.get("text", ""),
            "metadata": content.get("metadata", {})
        }
        
        # Füge strukturierte Daten hinzu
        if "text_blocks" in content:
            payload["text_blocks"] = content["text_blocks"]
        
        # Füge Detektionen hinzu, falls vorhanden und gewünscht
        if "detections" in content and options.get("include_detections", True):
            payload["detections"] = content["detections"]
        
        # API-Aufruf-Optionen
        headers = options.get("api_headers", {"Content-Type": "application/json"})
        timeout = options.get("api_timeout", 30)
        
        try:
            # Sende Daten an API
            response = requests.post(api_endpoint, json=payload, headers=headers, timeout=timeout)
            
            # Prüfe auf erfolgreichen Aufruf
            response.raise_for_status()
            
            # Speichere API-Antwort als Referenz
            response_data = response.json() if response.headers.get('content-type') == 'application/json' else {"status": "success"}
            
            # Optional, speichere API-Antwort in Datei
            if options.get("save_response", True):
                response_path = output_path.replace('.dat', '.json')
                with open(response_path, 'w', encoding='utf-8') as f:
                    json.dump(response_data, f, ensure_ascii=False, indent=2)
            
            return {
                "api_endpoint": api_endpoint,
                "status_code": response.status_code,
                "response": response_data,
                "response_path": response_path if options.get("save_response", True) else None
            }
        
        except requests.RequestException as e:
            self.logger.error(f"API-Aufruf fehlgeschlagen: {e}")
            return {
                "api_endpoint": api_endpoint,
                "error": str(e),
                "status": "failed"
            }
    
    def _generate_csv_output(self, content: Dict[str, Any], output_path: str,
                            options: Dict[str, Any]) -> Dict[str, Any]:
        """Generiert CSV-Ausgabe mit strukturierten Daten.
        
        Args:
            content: Verarbeiteter Inhalt
            output_path: Pfad für die Ausgabedatei
            options: Verarbeitungsoptionen
            
        Returns:
            Ergebnisinformationen
        """
        import csv
        
        # Überprüfe, ob strukturierte Daten vorhanden sind
        has_text_blocks = "text_blocks" in content and content["text_blocks"]
        has_detections = "detections" in content and content["detections"]
        
        if not (has_text_blocks or has_detections):
            self.logger.warning("Keine strukturierten Daten für CSV-Export")
            # Erstelle einfache CSV mit erkanntem Text
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(["document_id", "text"])
                writer.writerow([content.get("doc_id", ""), content.get("text", "")])
            
            return {
                "path": output_path,
                "format": "csv",
                "size": os.path.getsize(output_path),
                "rows": 1
            }
        
        # Bestimme den CSV-Exporttyp
        export_type = options.get("csv_export_type", "text_blocks" if has_text_blocks else "detections")
        
        row_count = 0
        if export_type == "text_blocks" and has_text_blocks:
            # Exportiere Textblöcke als CSV
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                # Header
                writer.writerow(["block_id", "text", "confidence", "x", "y", "width", "height"])
                
                # Daten
                for i, block in enumerate(content["text_blocks"]):
                    text = block.get("text", "")
                    confidence = block.get("confidence", 0.0)
                    bbox = block.get("bbox", [0, 0, 0, 0])
                    
                    writer.writerow([
                        i+1,
                        text,
                        confidence,
                        bbox[0], bbox[1], bbox[2], bbox[3]
                    ])
                    row_count += 1
        
        elif export_type == "detections" and has_detections:
            # Exportiere Objekterkennungen als CSV
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                # Header
                writer.writerow(["detection_id", "class_name", "confidence", "x", "y", "width", "height"])
                
                # Daten
                for i, detection in enumerate(content["detections"]):
                    class_name = detection.get("class_name", "unknown")
                    confidence = detection.get("confidence", 0.0)
                    box = detection.get("box", {"x": 0, "y": 0, "width": 0, "height": 0})
                    
                    writer.writerow([
                        i+1,
                        class_name,
                        confidence,
                        box["x"], box["y"], box["width"], box["height"]
                    ])
                    row_count += 1
        
        return {
            "path": output_path,
            "format": "csv",
            "size": os.path.getsize(output_path),
            "rows": row_count,
            "export_type": export_type
        }
    
    def _generate_markdown_output(self, content: Dict[str, Any], output_path: str,
                                options: Dict[str, Any]) -> Dict[str, Any]:
        """Generiert Markdown-Ausgabe mit formatiertem Text und Bildreferenzen.
        
        Args:
            content: Verarbeiteter Inhalt
            output_path: Pfad für die Ausgabedatei
            options: Verarbeitungsoptionen
            
        Returns:
            Ergebnisinformationen
        """
        # Initialisiere Markdown-Dokument
        md_lines = []
        
        # Titel
        doc_id = content.get("doc_id", "OCR-Ergebnis")
        md_lines.append(f"# {doc_id}")
        md_lines.append("")
        
        # Zeitstempel
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        md_lines.append(f"Generiert am: {timestamp}")
        md_lines.append("")
        
        # Füge Bild hinzu, wenn vorhanden und Bildausgabe existiert
        images_dir = options.get("images_dir", "")
        if "images" in content and options.get("include_images", True):
            # Prüfe, ob ein annotiertes Bild generiert wurde
            if "outputs" in options and "image_annotated" in options["outputs"]:
                image_path = options["outputs"]["image_annotated"].get("path", "")
                if image_path:
                    # Verwende relativen Pfad für Markdown
                    rel_path = os.path.relpath(image_path, os.path.dirname(output_path))
                    md_lines.append(f"![Annotiertes Bild]({rel_path})")
                    md_lines.append("")
        
        # Erkannter Text
        if "text" in content and content["text"]:
            md_lines.append("## Erkannter Text")
            md_lines.append("")
            md_lines.append(content["text"])
            md_lines.append("")
        
        # Metadaten
        if "metadata" in content and options.get("include_metadata", True):
            md_lines.append("## Metadaten")
            md_lines.append("")
            md_lines.append("| Attribut | Wert |")
            md_lines.append("|----------|------|")
            
            for key, value in content["metadata"].items():
                md_lines.append(f"| {key} | {value} |")
            
            md_lines.append("")
        
        # Erkannte Objekte
        if "detections" in content and options.get("include_detections", True):
            md_lines.append("## Erkannte Objekte")
            md_lines.append("")
            md_lines.append("| Klasse | Konfidenz | Position |")
            md_lines.append("|--------|-----------|----------|")
            
            for detection in content["detections"]:
                class_name = detection.get("class_name", "unbekannt")
                confidence = detection.get("confidence", 0.0)
                box = detection.get("box", {})
                position = f"x={box.get('x', 0)}, y={box.get('y', 0)}, w={box.get('width', 0)}, h={box.get('height', 0)}"
                
                md_lines.append(f"| {class_name} | {confidence:.2f} | {position} |")
            
            md_lines.append("")
        
        # Speichere Markdown-Datei
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(md_lines))
        
        return {
            "path": output_path,
            "format": "markdown",
            "size": os.path.getsize(output_path),
            "sections": md_lines.count("## ")
        }
    
    def _generate_html_output(self, content: Dict[str, Any], output_path: str,
                            options: Dict[str, Any]) -> Dict[str, Any]:
        """Generiert HTML-Ausgabe mit formatierten Ergebnissen und eingebetteten Bildern.
        
        Args:
            content: Verarbeiteter Inhalt
            output_path: Pfad für die Ausgabedatei
            options: Verarbeitungsoptionen
            
        Returns:
            Ergebnisinformationen
        """
        # Initialisiere HTML-Dokument
        html_lines = []
        html_lines.append("<!DOCTYPE html>")
        html_lines.append("<html lang='de'>")
        html_lines.append("<head>")
        html_lines.append("  <meta charset='UTF-8'>")
        html_lines.append("  <meta name='viewport' content='width=device-width, initial-scale=1.0'>")
        html_lines.append(f"  <title>{content.get('doc_id', 'OCR-Ergebnis')}</title>")
        html_lines.append("  <style>")
        html_lines.append("    body { font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }")
        html_lines.append("    .container { max-width: 1200px; margin: 0 auto; }")
        html_lines.append("    .header { background-color: #f5f5f5; padding: 10px; border-bottom: 1px solid #ddd; }")
        html_lines.append("    .result-image { max-width: 100%; margin: 20px 0; }")
        html_lines.append("    .text-content { white-space: pre-wrap; background-color: #f9f9f9; padding: 15px; border: 1px solid #ddd; }")
        html_lines.append("    table { border-collapse: collapse; width: 100%; }")
        html_lines.append("    th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }")
        html_lines.append("    th { background-color: #f2f2f2; }")
        html_lines.append("  </style>")
        html_lines.append("</head>")
        html_lines.append("<body>")
        html_lines.append("  <div class='container'>")
        
        # Header
        html_lines.append("    <div class='header'>")
        html_lines.append(f"      <h1>{content.get('doc_id', 'OCR-Ergebnis')}</h1>")
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        html_lines.append(f"      <p>Generiert am: {timestamp}</p>")
        html_lines.append("    </div>")
        
        # Bild, falls vorhanden
        if "images" in content and options.get("include_images", True):
            # Prüfe auf annotiertes Bild oder Originalbild
            image_path = None
            image_type = "Original"
            
            if "outputs" in options and "image_annotated" in options["outputs"]:
                image_path = options["outputs"]["image_annotated"].get("path", "")
                image_type = "Annotiert"
            
            if not image_path and "original" in content["images"]:
                original_image = content["images"]["original"]
                if isinstance(original_image, str) and os.path.exists(original_image):
                    image_path = original_image
            
            if image_path:
                # Erstelle Verzeichnis für Bilder relativ zur HTML-Datei
                html_dir = os.path.dirname(output_path)
                images_dir = os.path.join(html_dir, "images")
                os.makedirs(images_dir, exist_ok=True)
                
                # Kopiere oder konvertiere Bild
                image_filename = f"{os.path.splitext(os.path.basename(output_path))[0]}_image.jpg"
                html_image_path = os.path.join(images_dir, image_filename)
                
                # Kopiere oder konvertiere das Bild
                if isinstance(image_path, str) and os.path.exists(image_path):
                    img = Image.open(image_path)
                    img.save(html_image_path, "JPEG")
                elif isinstance(content["images"]["original"], np.ndarray):
                    cv2.imwrite(html_image_path, content["images"]["original"])
                
                # Füge Bild zur HTML hinzu
                rel_path = os.path.relpath(html_image_path, html_dir)
                html_lines.append(f"    <h2>{image_type}bild</h2>")
                html_lines.append(f"    <img src='{rel_path}' alt='{image_type}bild' class='result-image'>")
        
        # Erkannter Text
        if "text" in content and content["text"]:
            html_lines.append("    <h2>Erkannter Text</h2>")
            html_lines.append(f"    <div class='text-content'>{content['text']}</div>")
        
        # Erkannte Objekte
        if "detections" in content and options.get("include_detections", True) and content["detections"]:
            html_lines.append("    <h2>Erkannte Objekte</h2>")
            html_lines.append("    <table>")
            html_lines.append("      <tr><th>Klasse</th><th>Konfidenz</th><th>Position</th></tr>")
            
            for detection in content["detections"]:
                class_name = detection.get("class_name", "unbekannt")
                confidence = detection.get("confidence", 0.0)
                box = detection.get("box", {})
                position = f"x={box.get('x', 0)}, y={box.get('y', 0)}, w={box.get('width', 0)}, h={box.get('height', 0)}"
                
                html_lines.append(f"      <tr><td>{class_name}</td><td>{confidence:.2f}</td><td>{position}</td></tr>")
            
            html_lines.append("    </table>")
        
        # Metadaten
        if "metadata" in content and options.get("include_metadata", True) and content["metadata"]:
            html_lines.append("    <h2>Metadaten</h2>")
            html_lines.append("    <table>")
            html_lines.append("      <tr><th>Attribut</th><th>Wert</th></tr>")
            
            for key, value in content["metadata"].items():
                html_lines.append(f"      <tr><td>{key}</td><td>{value}</td></tr>")
            
            html_lines.append("    </table>")
        
        # Abschluss des HTML-Dokuments
        html_lines.append("  </div>")
        html_lines.append("</body>")
        html_lines.append("</html>")
        
        # Speichere HTML-Datei
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(html_lines))
        
        return {
            "path": output_path,
            "format": "html",
            "size": os.path.getsize(output_path),
            "includes_images": options.get("include_images", True) and "images" in content
        }
    
    def _generate_docx_output(self, content: Dict[str, Any], output_path: str,
                            options: Dict[str, Any]) -> Dict[str, Any]:
        """Generiert DOCX-Ausgabe für Microsoft Word-kompatible Dokumente.
        
        Args:
            content: Verarbeiteter Inhalt
            output_path: Pfad für die Ausgabedatei
            options: Verarbeitungsoptionen
            
        Returns:
            Ergebnisinformationen
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
        except ImportError:
            self.logger.error("python-docx nicht installiert. DOCX-Generierung nicht möglich.")
            raise ImportError("python-docx wird für DOCX-Generierung benötigt. Installieren Sie 'python-docx'.")
        
        # Erstelle neues DOCX-Dokument
        doc = Document()
        
        # Füge Titel hinzu
        doc.add_heading(content.get("doc_id", "OCR-Ergebnis"), level=1)
        
        # Zeitstempel
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Generiert am: {timestamp}")
        
        # Füge Bild hinzu, wenn vorhanden
        if "images" in content and options.get("include_images", True):
            doc.add_heading("Bild", level=2)
            
            # Prüfe auf annotiertes Bild oder Originalbild
            image_path = None
            
            if "outputs" in options and "image_annotated" in options["outputs"]:
                image_path = options["outputs"]["image_annotated"].get("path", "")
            
            if not image_path and "original" in content["images"]:
                original_image = content["images"]["original"]
                if isinstance(original_image, str) and os.path.exists(original_image):
                    image_path = original_image
                elif isinstance(original_image, np.ndarray):
                    # Speichere temporäres Bild
                    with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                        temp_filename = temp_file.name
                        cv2.imwrite(temp_filename, original_image)
                        image_path = temp_filename
            
            if image_path:
                try:
                    doc.add_picture(image_path, width=Inches(6))
                except Exception as e:
                    self.logger.error(f"Fehler beim Hinzufügen des Bildes zum DOCX: {e}")
                
                # Lösche temporäre Datei, falls erstellt
                if 'temp_filename' in locals() and os.path.exists(temp_filename):
                    try:
                        os.unlink(temp_filename)
                    except:
                        pass
        
        # Füge erkannten Text hinzu
        if "text" in content and content["text"]:
            doc.add_heading("Erkannter Text", level=2)
            doc.add_paragraph(content["text"])
        
        # Erkannte Objekte
        if "detections" in content and options.get("include_detections", True) and content["detections"]:
            doc.add_heading("Erkannte Objekte", level=2)
            
            # Erstelle Tabelle
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Füge Kopfzeile hinzu
            header_cells = table.rows[0].cells
            header_cells[0].text = "Klasse"
            header_cells[1].text = "Konfidenz"
            header_cells[2].text = "Position"
            
            # Füge Datenzeilen hinzu
            for detection in content["detections"]:
                class_name = detection.get("class_name", "unbekannt")
                confidence = detection.get("confidence", 0.0)
                box = detection.get("box", {})
                position = f"x={box.get('x', 0)}, y={box.get('y', 0)}, w={box.get('width', 0)}, h={box.get('height', 0)}"
                
                row_cells = table.add_row().cells
                row_cells[0].text = class_name
                row_cells[1].text = f"{confidence:.2f}"
                row_cells[2].text = position
        
        # Metadaten
        if "metadata" in content and options.get("include_metadata", True) and content["metadata"]:
            doc.add_heading("Metadaten", level=2)
            
            # Erstelle Tabelle
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Füge Kopfzeile hinzu
            header_cells = table.rows[0].cells
            header_cells[0].text = "Attribut"
            header_cells[1].text = "Wert"
            
            # Füge Datenzeilen hinzu
            for key, value in content["metadata"].items():
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = str(value)
        
        # Speichere DOCX-Datei
        doc.save(output_path)
        
        return {
            "path": output_path,
            "format": "docx",
            "size": os.path.getsize(output_path),
            "sections": len(doc.sections)
        }
    
    def _generate_combined_output(self, content: Dict[str, Any], output_path: str,
                                options: Dict[str, Any]) -> Dict[str, Any]:
        """Generiert ein kombiniertes Dokument mit allen verfügbaren Informationen.
        
        Dies erstellt ein umfassendes PDF-Dokument, das Text, Bilder, 
        Erkennungsergebnisse und andere Informationen kombiniert.
        
        Args:
            content: Verarbeiteter Inhalt
            output_path: Pfad für die Ausgabedatei
            options: Verarbeitungsoptionen
            
        Returns:
            Ergebnisinformationen
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
            from reportlab.platypus import Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
        except ImportError:
            self.logger.error("ReportLab nicht installiert. PDF-Generierung nicht möglich.")
            raise ImportError("ReportLab wird für PDF-Generierung benötigt. Installieren Sie 'reportlab'.")
        
        # Erstelle PDF-Dokument
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Benutzerdefinierte Stile
        styles.add(ParagraphStyle(
            name='OCRText',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
        ))
        
        styles.add(ParagraphStyle(
            name='Caption',
            parent=styles['Normal'],
            fontSize=10,
            leading=12,
            alignment=1,  # Zentriert
            spaceAfter=6,
        ))
        
        # Elemente für das PDF
        elements = []
        
        # Titel und Metadaten
        title = content.get("doc_id", "OCR-Ergebnis")
        elements.append(Paragraph(f"<b>{title}</b>", styles["Title"]))
        elements.append(Spacer(1, 0.25*inch))
        
        # Zeitstempel
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        elements.append(Paragraph(f"Generiert am: {timestamp}", styles["Normal"]))
        elements.append(Spacer(1, 0.25*inch))
        
        # Inhaltsübersicht (als Tabelle)
        overview_data = []
        overview_data.append(["Abschnitt", "Details"])
        
        # Füge Einträge basierend auf vorhandenen Inhalten hinzu
        if "text" in content and content["text"]:
            overview_data.append(["Text", f"{len(content['text'])} Zeichen"])
        
        if "text_blocks" in content:
            overview_data.append(["Textblöcke", f"{len(content['text_blocks'])} Blöcke"])
        
        if "detections" in content:
            overview_data.append(["Erkannte Objekte", f"{len(content['detections'])} Objekte"])
        
        if "metadata" in content:
            overview_data.append(["Metadaten", f"{len(content['metadata'])} Attribute"])
        
        if overview_data:
            table = Table(overview_data, colWidths=[1.5*inch, 4*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            elements.append(table)
            elements.append(Spacer(1, 0.25*inch))
        
        # Füge Bild hinzu, wenn vorhanden
        if "images" in content and options.get("include_images", True):
            elements.append(Paragraph("<b>Bild</b>", styles["Heading2"]))
            
            # Wähle annotiertes Bild oder Originalbild
            image_path = None
            image_type = "Original"
            
            if "outputs" in options and "image_annotated" in options["outputs"]:
                image_path = options["outputs"]["image_annotated"].get("path", "")
                image_type = "Annotiert"
            
            if not image_path and "original" in content["images"]:
                original_image = content["images"]["original"]
                if isinstance(original_image, str) and os.path.exists(original_image):
                    image_path = original_image
                elif isinstance(original_image, np.ndarray):
                    # Speichere temporäres Bild
                    with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                        temp_filename = temp_file.name
                        cv2.imwrite(temp_filename, original_image)
                        image_path = temp_filename
            
            if image_path:
                try:
                    img = RLImage(image_path, width=6*inch, height=4*inch, kind='proportional')
                    elements.append(img)
                    elements.append(Paragraph(f"{image_type}bild", styles["Caption"]))
                    elements.append(Spacer(1, 0.2*inch))
                except Exception as e:
                    self.logger.error(f"Fehler beim Hinzufügen des Bildes zum PDF: {e}")
                
                # Lösche temporäre Datei, falls erstellt
                if 'temp_filename' in locals() and os.path.exists(temp_filename):
                    try:
                        os.unlink(temp_filename)
                    except:
                        pass
        
        # Füge Thumbnail hinzu, falls vorhanden
        if "outputs" in options and "thumbnail" in options["outputs"]:
            thumbnail_path = options["outputs"]["thumbnail"].get("path", "")
            if thumbnail_path and os.path.exists(thumbnail_path):
                elements.append(Paragraph("<b>Thumbnail</b>", styles["Heading3"]))
                img = RLImage(thumbnail_path, width=2*inch, height=2*inch, kind='proportional')
                elements.append(img)
                elements.append(Spacer(1, 0.2*inch))
        
        # Füge erkannten Text hinzu
        if "text" in content and content["text"]:
            elements.append(Paragraph("<b>Erkannter Text</b>", styles["Heading2"]))
            elements.append(Spacer(1, 0.1*inch))
            
            text = content["text"]
            # Teile Text in Absätze
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    elements.append(Paragraph(para.replace('\n', '<br/>'), styles["OCRText"]))
                    elements.append(Spacer(1, 0.1*inch))
        
        # Textblöcke
        if "text_blocks" in content and content["text_blocks"] and options.get("include_text_blocks", True):
            elements.append(Paragraph("<b>Textblöcke</b>", styles["Heading2"]))
            elements.append(Spacer(1, 0.1*inch))
            
            # Erstelle Tabelle für Textblöcke
            block_data = [["ID", "Text", "Konfidenz", "Position"]]
            
            for i, block in enumerate(content["text_blocks"]):
                text = block.get("text", "")
                # Begrenze Text auf 50 Zeichen für Tabelle
                if len(text) > 50:
                    text = text[:47] + "..."
                
                confidence = block.get("confidence", 0.0)
                bbox = block.get("bbox", [0, 0, 0, 0])
                position = f"x={bbox[0]}, y={bbox[1]}, w={bbox[2]}, h={bbox[3]}"
                
                block_data.append([i+1, text, f"{confidence:.2f}", position])
            
            if len(block_data) > 1:
                table = Table(block_data, colWidths=[0.5*inch, 3*inch, 0.8*inch, 1.2*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                elements.append(table)
                elements.append(Spacer(1, 0.2*inch))
        
        # Erkannte Objekte
        if "detections" in content and content["detections"] and options.get("include_detections", True):
            elements.append(Paragraph("<b>Erkannte Objekte</b>", styles["Heading2"]))
            elements.append(Spacer(1, 0.1*inch))
            
            # Erstelle Tabelle für erkannte Objekte
            detection_data = [["ID", "Klasse", "Konfidenz", "Position"]]
            
            for i, detection in enumerate(content["detections"]):
                class_name = detection.get("class_name", "unbekannt")
                confidence = detection.get("confidence", 0.0)
                box = detection.get("box", {})
                position = f"x={box.get('x', 0)}, y={box.get('y', 0)}, w={box.get('width', 0)}, h={box.get('height', 0)}"
                
                detection_data.append([i+1, class_name, f"{confidence:.2f}", position])
            
            if len(detection_data) > 1:
                table = Table(detection_data, colWidths=[0.5*inch, 1.5*inch, 0.8*inch, 2.7*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                elements.append(table)
                elements.append(Spacer(1, 0.2*inch))
        
        # Metadaten
        if "metadata" in content and content["metadata"] and options.get("include_metadata", True):
            elements.append(Paragraph("<b>Metadaten</b>", styles["Heading2"]))
            elements.append(Spacer(1, 0.1*inch))
            
            # Erstelle Tabelle für Metadaten
            metadata_data = [["Attribut", "Wert"]]
            
            for key, value in content["metadata"].items():
                metadata_data.append([key, str(value)])
            
            if len(metadata_data) > 1:
                table = Table(metadata_data, colWidths=[2*inch, 4*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                elements.append(table)
                elements.append(Spacer(1, 0.2*inch))
        
        # Füge Verarbeitungsinformationen hinzu
        elements.append(Paragraph("<b>Verarbeitungsinformationen</b>", styles["Heading2"]))
        processing_data = [["Parameter", "Wert"]]
        
        # Füge verfügbare Verarbeitungsinformationen hinzu
        if "processing_info" in content:
            for key, value in content["processing_info"].items():
                processing_data.append([key, str(value)])
        
        processing_data.append(["Erstellungszeitpunkt", timestamp])
        processing_data.append(["Ausgabeformate", ", ".join(options.get("formats", self.default_formats))])
        
        if len(processing_data) > 1:
            table = Table(processing_data, colWidths=[2*inch, 4*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            elements.append(table)
        
        # Erstelle das PDF
        doc.build(elements)
        
        return {
            "path": output_path,
            "format": "combined",
            "size": os.path.getsize(output_path),
            "content_types": list(content.keys())
        }